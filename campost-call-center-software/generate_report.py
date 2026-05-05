#!/usr/bin/env python3
"""Generate Campost Call Center Software Proposal Report in Word format."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Style setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    h.font.name = 'Calibri'
    if i == 1:
        h.font.size = Pt(22)
        h.paragraph_format.space_before = Pt(24)
    elif i == 2:
        h.font.size = Pt(16)
        h.paragraph_format.space_before = Pt(18)
    else:
        h.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(12)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.style = doc.styles['Normal']
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(text)
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)


# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('PROPOSITION TECHNIQUE')
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Solution Logicielle pour le Centre d\'Appels\nde la CAMPOST')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

doc.add_paragraph()

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run('TECHNICAL PROPOSAL\nCall Center Software Solution for CAMPOST')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
run.italic = True

for _ in range(4):
    doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f'Date: {datetime.date.today().strftime("%B %d, %Y")}\nConfidential / Confidentiel')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading('Table of Contents / Table des Matieres', level=1)
toc_items = [
    '1. Executive Summary / Resume Executif',
    '2. Understanding of Campost / Comprehension de la CAMPOST',
    '3. Current Digital Ecosystem / Ecosysteme Numerique Actuel',
    '4. Call Center Needs Analysis / Analyse des Besoins',
    '5. Proposed Solution Architecture / Architecture Proposee',
    '6. Solution Comparison / Comparaison des Solutions',
    '7. Recommended Approach / Approche Recommandee',
    '8. Implementation Roadmap / Feuille de Route',
    '9. Compliance & Security / Conformite & Securite',
    '10. Cost Estimation / Estimation des Couts',
    '11. Why Choose Us / Pourquoi Nous Choisir',
    '12. Conclusion',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    p.runs[0].font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'This document presents our technical proposal in response to CAMPOST\'s '
    'appel d\'offres for the development and deployment of a Call Center Software '
    'Solution. CAMPOST (Societe des Postes du Cameroun), as Cameroon\'s national '
    'postal operator, serves millions of citizens through 234 post offices nationwide '
    'and an expanding portfolio of digital services including CampostMoney, the NPSI '
    '#237# USSD platform, and the upcoming Bolamba e-commerce platform.'
)
doc.add_paragraph(
    'Our research identifies that CAMPOST currently lacks a dedicated, integrated '
    'call center system. Customer service is handled through basic phone lines and '
    'email, which is insufficient for the scale and complexity of services offered. '
    'A modern call center solution will dramatically improve customer satisfaction, '
    'reduce response times, and support CAMPOST\'s digital transformation journey.'
)
doc.add_paragraph(
    'After evaluating 21 solutions across enterprise cloud, mid-market, Africa-focused, '
    'open-source, and custom-built categories, we recommend a hybrid approach combining '
    'proven open-source/Africa-focused technology with custom integrations for CAMPOST\'s '
    'unique ecosystem (CampostMoney, NPSI #237#, UPU IPS tracking, and Bolamba). This '
    'approach balances cost-effectiveness, local infrastructure compatibility, and '
    'compliance with Cameroon\'s new data protection law (Law No. 2024/017).'
)

doc.add_heading('Resume Executif (FR)', level=2)
doc.add_paragraph(
    'Ce document presente notre proposition technique en reponse a l\'appel d\'offres '
    'de la CAMPOST pour le developpement d\'une solution logicielle de centre d\'appels. '
    'Apres evaluation de 21 solutions, nous recommandons une approche hybride combinant '
    'des technologies open-source/africaines eprouvees avec des integrations personnalisees '
    'pour l\'ecosysteme unique de la CAMPOST. Cette approche optimise les couts tout en '
    'garantissant la conformite avec la loi camerounaise sur la protection des donnees '
    'personnelles (Loi No. 2024/017).'
)

doc.add_page_break()

# ============================================================
# 2. UNDERSTANDING OF CAMPOST
# ============================================================
doc.add_heading('2. Understanding of CAMPOST', level=1)
doc.add_paragraph(
    'CAMPOST is Cameroon\'s national postal operator, playing a critical role in the '
    'country\'s communication, logistics, and financial inclusion ecosystem.'
)

doc.add_heading('2.1 Core Services', level=2)
add_bullet(doc, 'Mail & Parcel Delivery: ', bold_prefix='')
add_bullet(doc, 'Domestic and international mail, EMS (Express Mail Service), parcel logistics')
add_bullet(doc, 'Financial Services (CampostMoney/CAMO): ', bold_prefix='')
add_bullet(doc, 'Savings accounts, loans, money transfers, bill payments (Camwater partnership)')
add_bullet(doc, 'Digital Platforms: ', bold_prefix='')
add_bullet(doc, 'NPSI #237# national USSD payment aggregation, Bolamba e-commerce (2026)')
add_bullet(doc, 'Infrastructure: ', bold_prefix='')
add_bullet(doc, '234 post offices nationwide, $52M data center in Yaounde')

doc.add_heading('2.2 Key Challenges for Customer Service', level=2)
add_bullet(doc, 'High volume of tracking inquiries across mail, parcels, and EMS')
add_bullet(doc, 'Financial service queries (CampostMoney account issues, transaction disputes)')
add_bullet(doc, 'Bilingual population requiring FR/EN service delivery')
add_bullet(doc, 'Variable internet connectivity across urban and rural post offices')
add_bullet(doc, 'No current integrated call center system (basic phone + email only)')
add_bullet(doc, 'Growing digital service portfolio increasing customer touchpoints')

doc.add_page_break()

# ============================================================
# 3. CURRENT DIGITAL ECOSYSTEM
# ============================================================
doc.add_heading('3. Current Digital Ecosystem', level=1)
doc.add_paragraph(
    'Understanding CAMPOST\'s existing infrastructure is essential for designing a '
    'call center solution that integrates seamlessly rather than operating in isolation.'
)

add_table(doc,
    ['Component', 'Description', 'Integration Priority'],
    [
        ['Data Center (Yaounde)', '$52M facility (AS328123), colocation services. Can host on-premise solutions.', 'HIGH - Hosting'],
        ['CampostMoney / CAMO', 'Digital banking platform by Sixgoldtech (London). Savings, loans, payments, transfers.', 'HIGH - Financial queries'],
        ['NPSI #237#', 'First African national USSD payment code. Ecobank, Afriland connected.', 'HIGH - Self-service channel'],
        ['Bolamba (2026)', 'E-commerce platform, CFA 2-3B investment, logistics hubs at Douala & Yaounde airports.', 'MEDIUM - Future integration'],
        ['FindMe (Nov 2024)', 'Digital addressing system via mobile app, partnership with FindMe.', 'LOW - Address lookup'],
        ['Camwater (Nov 2025)', 'Bill payment integration via CampostMoney.', 'MEDIUM - Payment queries'],
        ['UPU Membership', 'Member of Universal Postal Union, uses IPS for international tracking.', 'HIGH - Tracking queries'],
        ['234 Post Offices', 'Nationwide physical network, potential distributed agent locations.', 'HIGH - Multi-site'],
    ],
    col_widths=[4, 8, 4]
)

doc.add_page_break()

# ============================================================
# 4. CALL CENTER NEEDS ANALYSIS
# ============================================================
doc.add_heading('4. Call Center Needs Analysis', level=1)

doc.add_heading('4.1 Expected Call Volumes by Service', level=2)
add_table(doc,
    ['Service Area', 'Inquiry Type', 'Est. Volume', 'Complexity'],
    [
        ['Mail & Parcels', 'Package tracking, delivery status, complaints', 'Very High', 'Low-Medium'],
        ['CampostMoney', 'Account balance, transaction disputes, loan inquiries', 'High', 'Medium-High'],
        ['NPSI #237#', 'USSD transaction failures, onboarding, partner banks', 'Medium', 'Medium'],
        ['Bolamba (2026)', 'Order status, delivery issues, returns, seller support', 'Growing', 'Medium'],
        ['General', 'Post office hours, locations, postal rates, complaints', 'High', 'Low'],
        ['Bill Payments', 'Camwater and other bill payment confirmations', 'Medium', 'Low'],
    ]
)

doc.add_heading('4.2 Functional Requirements', level=2)
doc.add_heading('4.2.1 Core Call Center Features', level=3)
add_bullet(doc, 'Interactive Voice Response (IVR) with bilingual FR/EN menus')
add_bullet(doc, 'Automatic Call Distribution (ACD) with skills-based routing')
add_bullet(doc, 'Call queuing with estimated wait time announcements')
add_bullet(doc, 'Call recording for quality assurance and compliance')
add_bullet(doc, 'Callback scheduling to reduce hold times and telephony costs')
add_bullet(doc, 'Real-time supervisor dashboard with monitoring and whisper/barge')

doc.add_heading('4.2.2 Omnichannel Support', level=3)
add_bullet(doc, 'Voice (primary channel)')
add_bullet(doc, 'WhatsApp Business (dominant messaging app in Cameroon)')
add_bullet(doc, 'SMS for notifications and basic interactions')
add_bullet(doc, 'Email integration')
add_bullet(doc, 'USSD #237# self-service channel')
add_bullet(doc, 'Web chat on campost.cm and campostmoney.cm')

doc.add_heading('4.2.3 CAMPOST-Specific Integrations', level=3)
add_bullet(doc, 'UPU IPS: Real-time parcel/mail tracking within agent interface')
add_bullet(doc, 'CampostMoney API: Account lookup, transaction history, dispute management')
add_bullet(doc, 'NPSI #237#: USSD session handoff and self-service triggers')
add_bullet(doc, 'Bolamba: Order status, delivery tracking, return management')

doc.add_heading('4.2.4 Technical Requirements', level=3)
add_bullet(doc, 'Deployable on CAMPOST\'s Yaounde data center (on-premise or hybrid)')
add_bullet(doc, 'Operates on minimum 3-5 Mbps bandwidth for rural offices')
add_bullet(doc, 'GSM/SIP trunk support for MTN, Orange, and CAMTEL')
add_bullet(doc, 'Offline fallback mode for internet outages')
add_bullet(doc, 'Mobile agent app for lightweight deployment at post offices')
add_bullet(doc, 'Power outage resilience with graceful degradation')

doc.add_heading('4.2.5 Compliance Requirements', level=3)
add_bullet(doc, 'Cameroon Law No. 2024/017 on personal data protection (effective June 23, 2026)')
add_bullet(doc, 'COBAC Regulation No. 01/20 on consumer protection (CEMAC banking)')
add_bullet(doc, 'Data residency: customer data and recordings stored in Cameroon')
add_bullet(doc, 'FCFA (XAF) currency support in all financial reporting')

doc.add_page_break()

# ============================================================
# 5. PROPOSED SOLUTION ARCHITECTURE
# ============================================================
doc.add_heading('5. Proposed Solution Architecture', level=1)
doc.add_paragraph(
    'Based on our analysis of CAMPOST\'s needs, existing infrastructure, and the '
    'Cameroonian market context, we propose a modular, hybrid architecture that '
    'maximizes the use of CAMPOST\'s existing data center while leveraging proven '
    'technologies.'
)

doc.add_heading('5.1 Architecture Overview', level=2)
doc.add_paragraph(
    'The proposed solution follows a three-tier architecture deployed on CAMPOST\'s '
    'Yaounde data center:'
)

add_table(doc,
    ['Tier', 'Component', 'Technology'],
    [
        ['Communication Layer', 'PBX / Voice Gateway', 'Asterisk/Issabel + GSM Gateways (MTN/Orange/CAMTEL)'],
        ['Communication Layer', 'Omnichannel Router', 'WhatsApp Business API + SMS Gateway + Web Chat'],
        ['Application Layer', 'Agent Desktop', 'Custom web application (responsive, works on tablets)'],
        ['Application Layer', 'IVR Engine', 'Bilingual FR/EN auto-attendant with self-service'],
        ['Application Layer', 'Ticketing System', 'Integrated ticket management with SLA tracking'],
        ['Application Layer', 'Supervisor Dashboard', 'Real-time monitoring, reporting, quality management'],
        ['Integration Layer', 'CampostMoney Connector', 'API integration with Sixgoldtech platform'],
        ['Integration Layer', 'UPU IPS Connector', 'Parcel/mail tracking integration'],
        ['Integration Layer', 'NPSI #237# Bridge', 'USSD self-service and agent handoff'],
        ['Data Layer', 'Database', 'PostgreSQL (on-premise, CAMPOST data center)'],
        ['Data Layer', 'Call Recording Storage', 'Local NAS/SAN with encryption'],
        ['Data Layer', 'Analytics Engine', 'Business intelligence and reporting'],
    ],
    col_widths=[3.5, 4, 8]
)

doc.add_heading('5.2 Key Design Principles', level=2)
add_bullet(doc, 'On-premise first: ', bold_prefix='')
doc.add_paragraph('    All core components hosted on CAMPOST\'s data center for data sovereignty and compliance.')
add_bullet(doc, 'Low-bandwidth resilient: ', bold_prefix='')
doc.add_paragraph('    Optimized for 3-5 Mbps connections at rural post offices. Offline ticket mode available.')
add_bullet(doc, 'Mobile-ready agents: ', bold_prefix='')
doc.add_paragraph('    Responsive web app allows agents at post offices to use tablets instead of full desktop setups.')
add_bullet(doc, 'Bilingual by default: ', bold_prefix='')
doc.add_paragraph('    All interfaces, IVR menus, and reports available in French and English.')
add_bullet(doc, 'Modular integration: ', bold_prefix='')
doc.add_paragraph('    Each CAMPOST system (CampostMoney, NPSI, Bolamba) connected via standardized API connectors.')

doc.add_page_break()

# ============================================================
# 6. SOLUTION COMPARISON
# ============================================================
doc.add_heading('6. Solution Comparison', level=1)
doc.add_paragraph(
    'We evaluated 21 solutions across multiple categories. Below is a comparison '
    'of the most relevant options for CAMPOST\'s context.'
)

doc.add_heading('6.1 Enterprise Cloud Platforms', level=2)
add_table(doc,
    ['Solution', 'Strengths', 'Weaknesses for CAMPOST', 'Est. Annual Cost'],
    [
        ['Genesys Cloud CX', 'Market leader, full omnichannel, AI routing', 'No African data center, high cost, internet-dependent', '$150K-300K+'],
        ['Five9', 'Strong AI, workforce optimization', 'US-centric, no African presence, high bandwidth needs', '$120K-250K+'],
        ['NICE CXone', 'African data center (SA), enterprise-grade', 'High cost, complex deployment', '$130K-280K+'],
        ['Talkdesk', 'AI-powered, government deployments', 'No African presence, fully cloud-dependent', '$100K-200K+'],
    ],
    col_widths=[3, 4.5, 5, 3]
)
doc.add_paragraph('Assessment: Enterprise cloud platforms offer powerful features but are costly, internet-dependent, '
                   'and lack local presence in Cameroon. Not recommended as primary solution.', style='Normal')

doc.add_heading('6.2 Africa-Focused Solutions', level=2)
add_table(doc,
    ['Solution', 'Strengths', 'Weaknesses for CAMPOST', 'Est. Annual Cost'],
    [
        ['PressOne Africa', 'Works on 3 Mbps, Africa-native, proven scale', 'Nigeria-focused, limited francophone presence', '$20K-60K'],
        ['HelloDuty', 'Africa-built, zero-code IVR, API-rich', 'East Africa focus, no francophone markets yet', '$15K-50K'],
        ['Ameyo (Exotel)', 'Cloud + on-premise, French support, Africa presence', 'Limited Cameroon-specific support', '$30K-80K'],
        ['HoduCC', 'Affordable, Nigerian deployments, omnichannel', 'India-based support, timezone challenges', '$15K-45K'],
    ],
    col_widths=[3, 4.5, 5, 3]
)
doc.add_paragraph('Assessment: Africa-focused solutions offer better pricing and bandwidth optimization, '
                   'but none have established presence in Cameroon/Central Africa.', style='Normal')

doc.add_heading('6.3 Open Source / On-Premise', level=2)
add_table(doc,
    ['Solution', 'Strengths', 'Weaknesses for CAMPOST', 'Est. Cost (Setup + Annual)'],
    [
        ['Asterisk', 'Free, highly customizable, massive community', 'Requires deep expertise, no GUI by default', '$30K-60K (services)'],
        ['Issabel', 'Free, modern UI, French community, full PBX+CC', 'Smaller community than Asterisk', '$25K-50K (services)'],
        ['Vicidial', 'Free, proven call center features', 'Dated interface, outbound-focused', '$20K-45K (services)'],
        ['3CX', 'On-premise option, modern, easy to use', 'Limited call center features, license cost', '$10K-30K (license+setup)'],
    ],
    col_widths=[3, 4.5, 5, 3.5]
)
doc.add_paragraph('Assessment: Open-source solutions offer the best cost profile and can be deployed on '
                   'CAMPOST\'s data center. Require custom development for CAMPOST integrations.', style='Normal')

doc.add_heading('6.4 Specialized & Postal-Specific', level=2)
add_table(doc,
    ['Solution', 'Strengths', 'Weaknesses for CAMPOST', 'Est. Cost'],
    [
        ['UPU PTC Solutions', 'Native postal integration, IPS/IFS compatible', 'Limited call center features, supplement needed', 'UPU membership fees'],
        ['Odoo VoIP+Helpdesk', 'Full ERP, open-source, French support', 'VoIP module is basic, needs customization', '$20K-50K'],
        ['Contaque', 'GSM-based calling, AI features', 'Smaller vendor, limited Africa references', '$25K-60K'],
        ['Custom-built', 'Perfect fit, full control, data sovereignty', 'Higher upfront cost, longer timeline', '$80K-150K (build)'],
    ],
    col_widths=[3, 4.5, 5, 3]
)

doc.add_page_break()

# ============================================================
# 7. RECOMMENDED APPROACH
# ============================================================
doc.add_heading('7. Recommended Approach', level=1)
doc.add_paragraph(
    'Based on our comprehensive evaluation, we recommend the following approach, '
    'ranked by suitability for CAMPOST\'s specific context:'
)

doc.add_heading('7.1 Option A: Hybrid Custom Solution (Recommended)', level=2)
p = doc.add_paragraph()
run = p.add_run('Best balance of cost, control, compliance, and CAMPOST integration.')
run.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x6B, 0x1B)

doc.add_paragraph('Core stack:')
add_bullet(doc, 'Issabel/Asterisk: ', bold_prefix='')
doc.add_paragraph('    Open-source PBX and call center engine, deployed on CAMPOST data center')
add_bullet(doc, 'Custom Agent Desktop: ', bold_prefix='')
doc.add_paragraph('    Web-based application with integrated CampostMoney, UPU IPS, and NPSI views')
add_bullet(doc, 'WhatsApp Business API: ', bold_prefix='')
doc.add_paragraph('    Via CM.com or direct Meta integration for omnichannel')
add_bullet(doc, 'GSM Gateways: ', bold_prefix='')
doc.add_paragraph('    Direct connection to MTN, Orange, and CAMTEL networks')
add_bullet(doc, 'PostgreSQL + Local Storage: ', bold_prefix='')
doc.add_paragraph('    All data remains in CAMPOST\'s data center')

doc.add_paragraph()
add_table(doc,
    ['Aspect', 'Details'],
    [
        ['Estimated Budget', 'FCFA 50-90M ($80K-150K) for development + Year 1'],
        ['Annual Running Cost', 'FCFA 12-25M ($20K-40K) for maintenance, telephony, hosting'],
        ['Implementation Time', '4-6 months (phased rollout)'],
        ['Data Sovereignty', '100% on-premise at CAMPOST data center'],
        ['Scalability', 'Unlimited agents, grow as needed'],
        ['Compliance', 'Full control over Law No. 2024/017 compliance'],
    ],
    col_widths=[4, 12]
)

doc.add_heading('7.2 Option B: Africa-Focused Cloud + Custom Integration', level=2)
doc.add_paragraph('Use PressOne Africa or Ameyo as the call center engine, with custom API '
                   'connectors for CAMPOST systems.')
add_table(doc,
    ['Aspect', 'Details'],
    [
        ['Estimated Budget', 'FCFA 30-55M ($50K-90K) for setup + Year 1'],
        ['Annual Running Cost', 'FCFA 18-45M ($30K-75K) subscription + telephony'],
        ['Implementation Time', '2-4 months'],
        ['Data Sovereignty', 'Partial - voice/agent on cloud, integrations on-premise'],
        ['Risk', 'Vendor dependency, recurring subscription costs, data residency concerns'],
    ],
    col_widths=[4, 12]
)

doc.add_heading('7.3 Option C: Enterprise Cloud Platform', level=2)
doc.add_paragraph('Deploy Genesys Cloud CX or NICE CXone for maximum features.')
add_table(doc,
    ['Aspect', 'Details'],
    [
        ['Estimated Budget', 'FCFA 90-180M ($150K-300K) for Year 1'],
        ['Annual Running Cost', 'FCFA 75-170M ($120K-280K) subscription'],
        ['Implementation Time', '3-6 months'],
        ['Data Sovereignty', 'Limited - data in US/EU/South Africa'],
        ['Risk', 'Very high cost, full internet dependency, limited local support'],
    ],
    col_widths=[4, 12]
)

doc.add_page_break()

# ============================================================
# 8. IMPLEMENTATION ROADMAP
# ============================================================
doc.add_heading('8. Implementation Roadmap (Option A)', level=1)

add_table(doc,
    ['Phase', 'Timeline', 'Activities', 'Deliverables'],
    [
        ['Phase 1: Foundation', 'Weeks 1-4', 'Requirements validation, infrastructure setup on CAMPOST data center, Issabel/Asterisk installation, GSM gateway configuration', 'Working PBX, basic IVR (FR/EN), internal test calls'],
        ['Phase 2: Core Features', 'Weeks 5-10', 'Agent desktop development, ACD/routing rules, call recording, ticketing system, supervisor dashboard', 'Functional call center with 10-20 agents, basic reporting'],
        ['Phase 3: Integrations', 'Weeks 11-16', 'CampostMoney API integration, UPU IPS tracking integration, NPSI #237# USSD bridge, WhatsApp Business channel', 'Omnichannel center with CAMPOST system integrations'],
        ['Phase 4: Pilot', 'Weeks 17-20', 'Pilot at Yaounde HQ + 5 major post offices, agent training (FR/EN), performance tuning, user acceptance testing', 'Validated system, trained agents, performance benchmarks'],
        ['Phase 5: Rollout', 'Weeks 21-24', 'Phased rollout to remaining offices, remote agent setup, knowledge base population, go-live support', 'Full production deployment, documentation, handover'],
    ],
    col_widths=[3, 2.5, 5.5, 5]
)

doc.add_heading('8.1 Post-Deployment', level=2)
add_bullet(doc, 'Month 7-12: Bolamba e-commerce integration when platform launches')
add_bullet(doc, 'Ongoing: AI-powered IVR enhancement, sentiment analysis, predictive routing')
add_bullet(doc, 'Quarterly: Performance reviews, agent training refresh, system updates')

doc.add_page_break()

# ============================================================
# 9. COMPLIANCE & SECURITY
# ============================================================
doc.add_heading('9. Compliance & Security', level=1)

doc.add_heading('9.1 Cameroon Data Protection Law (No. 2024/017)', level=2)
doc.add_paragraph('Effective June 23, 2026, this law requires:')
add_bullet(doc, 'Data residency: Customer data stored within Cameroon (CAMPOST data center)')
add_bullet(doc, 'Consent management: Recording consent for call recordings')
add_bullet(doc, 'Breach notification: Incident response procedures')
add_bullet(doc, 'Data minimization: Collect only necessary customer information')
add_bullet(doc, 'Right to access/deletion: Customer data management portal')
doc.add_paragraph('Our recommended on-premise approach (Option A) provides full compliance by default.')

doc.add_heading('9.2 CEMAC Financial Regulations', level=2)
doc.add_paragraph('COBAC Regulation No. 01/20 requires:')
add_bullet(doc, 'Call recordings for financial service interactions (CampostMoney)')
add_bullet(doc, 'Audit trail for all customer dispute resolutions')
add_bullet(doc, 'Secure authentication before disclosing account information')
add_bullet(doc, 'Regular compliance reporting')

doc.add_heading('9.3 Security Measures', level=2)
add_bullet(doc, 'End-to-end encryption for voice and data channels')
add_bullet(doc, 'Role-based access control (RBAC) for agent/supervisor/admin roles')
add_bullet(doc, 'Encrypted call recording storage with retention policies')
add_bullet(doc, 'Regular security audits and penetration testing')
add_bullet(doc, 'Backup and disaster recovery on CAMPOST infrastructure')

doc.add_page_break()

# ============================================================
# 10. COST ESTIMATION
# ============================================================
doc.add_heading('10. Cost Estimation (Option A - Recommended)', level=1)

doc.add_heading('10.1 Initial Investment (Year 1)', level=2)
add_table(doc,
    ['Item', 'Description', 'Est. Cost (FCFA)', 'Est. Cost (USD)'],
    [
        ['Infrastructure Setup', 'Server config on CAMPOST DC, network, GSM gateways', '8,000,000', '$13,000'],
        ['Issabel/Asterisk Deploy', 'PBX installation, IVR config, SIP trunks', '5,000,000', '$8,000'],
        ['Agent Desktop Dev', 'Custom web app with CAMPOST integrations', '20,000,000', '$33,000'],
        ['CampostMoney Integration', 'API connector for financial service queries', '8,000,000', '$13,000'],
        ['UPU IPS Integration', 'Tracking system connector', '5,000,000', '$8,000'],
        ['NPSI #237# Bridge', 'USSD self-service integration', '6,000,000', '$10,000'],
        ['WhatsApp Channel', 'WhatsApp Business API setup + integration', '4,000,000', '$6,500'],
        ['Training & Documentation', 'Agent training (FR/EN), admin guides, videos', '4,000,000', '$6,500'],
        ['GSM Gateway Hardware', '4-8 port gateways for MTN/Orange/CAMTEL', '3,000,000', '$5,000'],
        ['Testing & QA', 'UAT, load testing, security audit', '3,000,000', '$5,000'],
        ['Project Management', 'PM, coordination, reporting (6 months)', '6,000,000', '$10,000'],
        ['', 'TOTAL INITIAL INVESTMENT', '72,000,000', '$118,000'],
    ],
    col_widths=[3.5, 5.5, 3.5, 3]
)

doc.add_heading('10.2 Annual Recurring Costs', level=2)
add_table(doc,
    ['Item', 'Description', 'Est. Annual (FCFA)', 'Est. Annual (USD)'],
    [
        ['Telephony (SIP/GSM)', 'Voice minutes, SMS, trunk fees', '6,000,000', '$10,000'],
        ['WhatsApp Business API', 'Message fees and API costs', '2,000,000', '$3,300'],
        ['Maintenance & Support', 'Bug fixes, updates, 8x5 support', '8,000,000', '$13,000'],
        ['Hosting (CAMPOST DC)', 'Power, rack space, bandwidth', '2,000,000', '$3,300'],
        ['Licenses (if any)', '3CX or additional modules', '1,500,000', '$2,500'],
        ['Training (ongoing)', 'New agent onboarding, refresher courses', '2,000,000', '$3,300'],
        ['', 'TOTAL ANNUAL RECURRING', '21,500,000', '$35,400'],
    ],
    col_widths=[3.5, 5.5, 3.5, 3]
)

doc.add_heading('10.3 3-Year Total Cost of Ownership', level=2)
add_table(doc,
    ['Period', 'Cost (FCFA)', 'Cost (USD)'],
    [
        ['Year 1 (Build + Run)', '93,500,000', '$153,400'],
        ['Year 2 (Run + Bolamba)', '28,000,000', '$46,000'],
        ['Year 3 (Run + Optimization)', '21,500,000', '$35,400'],
        ['3-Year TCO', '143,000,000', '$234,800'],
    ],
    col_widths=[5, 4, 4]
)

doc.add_page_break()

# ============================================================
# 11. WHY CHOOSE US
# ============================================================
doc.add_heading('11. Why Choose Us', level=1)
add_bullet(doc, 'Local Expertise: ', bold_prefix='')
doc.add_paragraph('    Deep understanding of Cameroon\'s telecom landscape (MTN, Orange, CAMTEL), '
                   'bilingual FR/EN market, and FCFA business environment.')
add_bullet(doc, 'CAMPOST Integration Knowledge: ', bold_prefix='')
doc.add_paragraph('    Thorough research of CAMPOST\'s ecosystem including CampostMoney, NPSI #237#, '
                   'UPU IPS, and upcoming Bolamba platform.')
add_bullet(doc, 'Compliance-First Approach: ', bold_prefix='')
doc.add_paragraph('    Solution designed from the ground up for Law No. 2024/017 and COBAC compliance.')
add_bullet(doc, 'Cost-Effective Technology: ', bold_prefix='')
doc.add_paragraph('    Open-source foundation eliminates expensive licensing while delivering enterprise features.')
add_bullet(doc, 'Data Sovereignty: ', bold_prefix='')
doc.add_paragraph('    100% on-premise deployment on CAMPOST\'s own data center - no data leaves Cameroon.')
add_bullet(doc, 'Scalable & Future-Ready: ', bold_prefix='')
doc.add_paragraph('    Modular architecture ready for Bolamba integration and AI enhancements.')

doc.add_page_break()

# ============================================================
# 12. CONCLUSION
# ============================================================
doc.add_heading('12. Conclusion', level=1)
doc.add_paragraph(
    'CAMPOST stands at a pivotal moment in its digital transformation. With CampostMoney '
    'serving a growing customer base, the NPSI #237# platform connecting banks nationwide, '
    'and the Bolamba e-commerce platform launching in 2026, a modern call center is no longer '
    'optional - it is essential infrastructure for customer satisfaction and service quality.'
)
doc.add_paragraph(
    'Our recommended hybrid approach (Option A) delivers a comprehensive, CAMPOST-integrated '
    'call center solution at a fraction of the cost of enterprise cloud platforms, while '
    'ensuring full data sovereignty and regulatory compliance. By leveraging CAMPOST\'s '
    'existing $52M data center and proven open-source technologies, we minimize recurring '
    'costs and maximize long-term value.'
)
doc.add_paragraph(
    'The 3-year total cost of ownership of approximately FCFA 143 million ($235K) compares '
    'favorably to enterprise alternatives that would cost FCFA 225-510 million ($370K-840K) '
    'over the same period, while offering superior integration with CAMPOST\'s unique ecosystem.'
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('We look forward to partnering with CAMPOST on this transformative project.')
run.bold = True
run.italic = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

doc.add_page_break()

# ============================================================
# APPENDIX: SOURCES
# ============================================================
doc.add_heading('Appendix: Research Sources', level=1)
sources = [
    'Campost Bets CFA3bn on Bolamba E-Commerce Platform - Business in Cameroon',
    'Campost + FindMe Digital Addressing Partnership - We Are Tech Africa',
    'CampostMoney Digital Banking Platform - campostmoney.com',
    'Sixgoldtech CampostMoney Implementation - sixgoldtech.com',
    'USSD #237# National Payment Platform Launch - Digital Business Africa',
    'Cameroon First African Country with National USSD Code - SUPTIC',
    'Camwater-Campost Bill Payment Partnership - Investir au Cameroun',
    'Campost Data Center (AS328123) - PeeringDB',
    'Cameroon Telecom Sector Grows 18% in 2024 - TechAfrica News',
    'Cameroon Telecom Operators Report 2025 - GlobeNewsWire',
    'Cameroon Data Protection Law No. 2024/017 - LEX Africa',
    'CEMAC Data Protection Framework - DataGuidance',
    'NICE CXone South Africa Instance (Dec 2025) - NICE',
    'Best Call Centre Solution Providers in Africa (2024) - Telebu',
    'PressOne Africa Call Center - pressone.africa',
    'HelloDuty AI-Powered Call Center - helloduty.com',
    'Ameyo Cloud Contact Center in Africa - ameyo.com',
    'HoduCC Contact Center Software - hodusoft.com',
    'UPU Postal Technology Centre - upu.int',
    'MINPOSTEL Government Modernization - minpostel.gov.cm',
]
for i, src in enumerate(sources, 1):
    p = doc.add_paragraph(f'{i}. {src}')
    p.runs[0].font.size = Pt(10)

# Save
output_path = '/Users/user/Documents/hairstyle/hairstyle/campost-call-center-software/Campost_Call_Center_Proposal.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
