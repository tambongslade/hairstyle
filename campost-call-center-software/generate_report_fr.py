#!/usr/bin/env python3
"""Generate Campost Call Center Software Proposal Report in French."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    h.font.name = 'Calibri'
    if i == 1:
        h.font.size = Pt(22)
        h.paragraph_format.space_before = Pt(24)
    elif i == 2:
        h.font.size = Pt(16)
        h.paragraph_format.space_before = Pt(18)
    else:
        h.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(12)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.style = doc.styles['Normal']
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(text)
    run.font.size = Pt(11)


# ============================================================
# PAGE DE COUVERTURE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('PROPOSITION TECHNIQUE')
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Solution Logicielle pour le Centre d\'Appels\nde la CAMPOST')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

doc.add_paragraph()

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run('Societe des Postes du Cameroun')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
run.italic = True

for _ in range(4):
    doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f'Date : {datetime.date.today().strftime("%d/%m/%Y")}\nDocument Confidentiel')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ============================================================
# TABLE DES MATIERES
# ============================================================
doc.add_heading('Table des Matieres', level=1)
toc_items = [
    '1. Resume Executif',
    '2. Presentation de la CAMPOST',
    '3. Ecosysteme Numerique Actuel',
    '4. Analyse des Besoins du Centre d\'Appels',
    '5. Architecture de la Solution Proposee',
    '6. Comparaison des Solutions',
    '7. Approche Recommandee',
    '8. Plan de Mise en Oeuvre',
    '9. Conformite et Securite',
    '10. Estimation des Couts',
    '11. Pourquoi Nous Choisir',
    '12. Conclusion',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    p.runs[0].font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 1. RESUME EXECUTIF
# ============================================================
doc.add_heading('1. Resume Executif', level=1)
doc.add_paragraph(
    'Ce document presente notre proposition technique en reponse a l\'appel d\'offres '
    'de la CAMPOST (Societe des Postes du Cameroun) pour le developpement et le '
    'deploiement d\'une solution logicielle de centre d\'appels. La CAMPOST, en tant '
    'qu\'operateur postal national du Cameroun, dessert des millions de citoyens a '
    'travers 234 bureaux de poste sur l\'ensemble du territoire national et un '
    'portefeuille croissant de services numeriques comprenant CampostMoney, la '
    'plateforme USSD NPSI #237# et la future plateforme e-commerce Bolamba.'
)
doc.add_paragraph(
    'Notre recherche revele que la CAMPOST ne dispose pas actuellement d\'un systeme '
    'de centre d\'appels dedie et integre. Le service client est assure par des lignes '
    'telephoniques basiques et le courrier electronique, ce qui est insuffisant pour '
    'l\'ampleur et la complexite des services offerts. Une solution moderne de centre '
    'd\'appels ameliorera considerablement la satisfaction client, reduira les delais '
    'de reponse et soutiendra la transformation numerique de la CAMPOST.'
)
doc.add_paragraph(
    'Apres avoir evalue 21 solutions a travers differentes categories (cloud entreprise, '
    'marche intermediaire, solutions africaines, open source et solutions sur mesure), '
    'nous recommandons une approche hybride combinant des technologies open source et '
    'africaines eprouvees avec des integrations personnalisees pour l\'ecosysteme unique '
    'de la CAMPOST (CampostMoney, NPSI #237#, suivi UPU IPS et Bolamba). Cette approche '
    'offre le meilleur equilibre entre cout, controle, conformite avec la nouvelle loi '
    'camerounaise sur la protection des donnees personnelles (Loi N. 2024/017) et '
    'compatibilite avec l\'infrastructure locale.'
)

doc.add_page_break()

# ============================================================
# 2. PRESENTATION DE LA CAMPOST
# ============================================================
doc.add_heading('2. Presentation de la CAMPOST', level=1)
doc.add_paragraph(
    'La CAMPOST est l\'operateur postal national du Cameroun, jouant un role essentiel '
    'dans l\'ecosysteme de communication, de logistique et d\'inclusion financiere du pays.'
)

doc.add_heading('2.1 Services Principaux', level=2)
add_bullet(doc, 'Courrier et Colis : Distribution nationale et internationale, EMS (Express Mail Service), logistique de colis')
add_bullet(doc, 'Services Financiers (CampostMoney/CAMO) : Comptes d\'epargne, prets, transferts d\'argent, paiements de factures (partenariat Camwater)')
add_bullet(doc, 'Plateformes Numeriques : Plateforme d\'agregation de paiements USSD NPSI #237#, e-commerce Bolamba (2026)')
add_bullet(doc, 'Infrastructure : 234 bureaux de poste a travers le pays, centre de donnees de 52 millions de dollars a Yaounde')

doc.add_heading('2.2 Defis Actuels du Service Client', level=2)
add_bullet(doc, 'Volume eleve de demandes de suivi de courrier, colis et EMS')
add_bullet(doc, 'Demandes liees aux services financiers (problemes de compte CampostMoney, litiges de transactions)')
add_bullet(doc, 'Population bilingue necessitant un service en francais et en anglais')
add_bullet(doc, 'Connectivite internet variable entre les bureaux de poste urbains et ruraux')
add_bullet(doc, 'Absence de systeme de centre d\'appels integre (telephone basique + email uniquement)')
add_bullet(doc, 'Portefeuille croissant de services numeriques augmentant les points de contact client')

doc.add_page_break()

# ============================================================
# 3. ECOSYSTEME NUMERIQUE ACTUEL
# ============================================================
doc.add_heading('3. Ecosysteme Numerique Actuel', level=1)
doc.add_paragraph(
    'La comprehension de l\'infrastructure existante de la CAMPOST est essentielle '
    'pour concevoir une solution de centre d\'appels qui s\'integre harmonieusement '
    'plutot que de fonctionner de maniere isolee.'
)

add_table(doc,
    ['Composant', 'Description', 'Priorite d\'Integration'],
    [
        ['Centre de Donnees (Yaounde)', 'Installation de 52M$ (AS328123), services de colocation. Peut heberger des solutions sur site.', 'HAUTE - Hebergement'],
        ['CampostMoney / CAMO', 'Plateforme bancaire numerique par Sixgoldtech (Londres). Epargne, prets, paiements, transferts.', 'HAUTE - Requetes financieres'],
        ['NPSI #237#', 'Premier code USSD national de paiement en Afrique. Ecobank, Afriland connectes.', 'HAUTE - Canal libre-service'],
        ['Bolamba (2026)', 'Plateforme e-commerce, investissement de 2-3 milliards FCFA, hubs logistiques a Douala et Yaounde.', 'MOYENNE - Integration future'],
        ['FindMe (Nov 2024)', 'Systeme d\'adressage numerique via application mobile.', 'BASSE - Recherche d\'adresses'],
        ['Camwater (Nov 2025)', 'Integration du paiement de factures via CampostMoney.', 'MOYENNE - Requetes de paiement'],
        ['Membre de l\'UPU', 'Membre de l\'Union Postale Universelle, utilise IPS pour le suivi international.', 'HAUTE - Requetes de suivi'],
        ['234 Bureaux de Poste', 'Reseau physique national, emplacements potentiels pour agents distribues.', 'HAUTE - Multi-sites'],
    ],
    col_widths=[4, 8, 4]
)

doc.add_page_break()

# ============================================================
# 4. ANALYSE DES BESOINS
# ============================================================
doc.add_heading('4. Analyse des Besoins du Centre d\'Appels', level=1)

doc.add_heading('4.1 Volumes d\'Appels Prevus par Service', level=2)
add_table(doc,
    ['Domaine de Service', 'Type de Demande', 'Volume Estime', 'Complexite'],
    [
        ['Courrier et Colis', 'Suivi de colis, statut de livraison, reclamations', 'Tres Eleve', 'Faible-Moyen'],
        ['CampostMoney', 'Solde de compte, litiges de transactions, demandes de prets', 'Eleve', 'Moyen-Eleve'],
        ['NPSI #237#', 'Echecs de transactions USSD, inscription, banques partenaires', 'Moyen', 'Moyen'],
        ['Bolamba (2026)', 'Statut de commande, problemes de livraison, retours', 'Croissant', 'Moyen'],
        ['General', 'Horaires des bureaux, localisations, tarifs postaux, plaintes', 'Eleve', 'Faible'],
        ['Paiements de Factures', 'Confirmations de paiement Camwater et autres', 'Moyen', 'Faible'],
    ]
)

doc.add_heading('4.2 Exigences Fonctionnelles', level=2)

doc.add_heading('4.2.1 Fonctionnalites de Base du Centre d\'Appels', level=3)
add_bullet(doc, 'Serveur Vocal Interactif (SVI) avec menus bilingues francais/anglais')
add_bullet(doc, 'Distribution Automatique des Appels (DAA) avec routage base sur les competences')
add_bullet(doc, 'File d\'attente avec annonce du temps d\'attente estime')
add_bullet(doc, 'Enregistrement des appels pour l\'assurance qualite et la conformite')
add_bullet(doc, 'Planification de rappels pour reduire les temps d\'attente et les couts de telephonie')
add_bullet(doc, 'Tableau de bord superviseur en temps reel avec ecoute et intervention')

doc.add_heading('4.2.2 Support Multicanal', level=3)
add_bullet(doc, 'Voix (canal principal)')
add_bullet(doc, 'WhatsApp Business (application de messagerie dominante au Cameroun)')
add_bullet(doc, 'SMS pour les notifications et interactions basiques')
add_bullet(doc, 'Integration email')
add_bullet(doc, 'Canal libre-service USSD #237#')
add_bullet(doc, 'Chat en ligne sur campost.cm et campostmoney.cm')

doc.add_heading('4.2.3 Integrations Specifiques CAMPOST', level=3)
add_bullet(doc, 'UPU IPS : Suivi en temps reel des colis et courriers dans l\'interface agent')
add_bullet(doc, 'API CampostMoney : Consultation de comptes, historique de transactions, gestion des litiges')
add_bullet(doc, 'NPSI #237# : Transfert de session USSD et declenchement du libre-service')
add_bullet(doc, 'Bolamba : Statut des commandes, suivi des livraisons, gestion des retours')

doc.add_heading('4.2.4 Exigences Techniques', level=3)
add_bullet(doc, 'Deployable sur le centre de donnees CAMPOST de Yaounde (sur site ou hybride)')
add_bullet(doc, 'Fonctionne avec un minimum de 3-5 Mbps de bande passante pour les bureaux ruraux')
add_bullet(doc, 'Support des lignes telephoniques GSM/SIP pour MTN, Orange et CAMTEL')
add_bullet(doc, 'Mode de secours hors ligne en cas de coupure internet')
add_bullet(doc, 'Application mobile pour agents, permettant l\'utilisation de tablettes dans les bureaux de poste')
add_bullet(doc, 'Resilience aux coupures de courant avec degradation progressive')

doc.add_heading('4.2.5 Exigences de Conformite', level=3)
add_bullet(doc, 'Loi camerounaise N. 2024/017 sur la protection des donnees personnelles (en vigueur le 23 juin 2026)')
add_bullet(doc, 'Reglement COBAC N. 01/20 sur la protection des consommateurs de services bancaires (CEMAC)')
add_bullet(doc, 'Residence des donnees : donnees clients et enregistrements stockes au Cameroun')
add_bullet(doc, 'Support de la devise FCFA (XAF) dans tous les rapports financiers')

doc.add_page_break()

# ============================================================
# 5. ARCHITECTURE DE LA SOLUTION
# ============================================================
doc.add_heading('5. Architecture de la Solution Proposee', level=1)
doc.add_paragraph(
    'Sur la base de notre analyse des besoins de la CAMPOST, de l\'infrastructure existante '
    'et du contexte camerounais, nous proposons une architecture modulaire et hybride qui '
    'maximise l\'utilisation du centre de donnees existant de la CAMPOST tout en tirant parti '
    'de technologies eprouvees.'
)

doc.add_heading('5.1 Vue d\'Ensemble de l\'Architecture', level=2)

add_table(doc,
    ['Couche', 'Composant', 'Technologie'],
    [
        ['Communication', 'Central Telephonique / Passerelle Vocale', 'Issabel/Asterisk + Passerelles GSM (MTN/Orange/CAMTEL)'],
        ['Communication', 'Routeur Multicanal', 'API WhatsApp Business + Passerelle SMS + Chat Web'],
        ['Application', 'Poste de Travail Agent', 'Application web personnalisee (responsive, fonctionne sur tablettes)'],
        ['Application', 'Moteur SVI', 'Serveur vocal bilingue FR/EN avec libre-service'],
        ['Application', 'Systeme de Tickets', 'Gestion integree des tickets avec suivi des delais'],
        ['Application', 'Tableau de Bord Superviseur', 'Surveillance en temps reel, rapports, gestion de la qualite'],
        ['Integration', 'Connecteur CampostMoney', 'Integration API avec la plateforme Sixgoldtech'],
        ['Integration', 'Connecteur UPU IPS', 'Integration du suivi des colis et courriers'],
        ['Integration', 'Pont NPSI #237#', 'Libre-service USSD et transfert vers agent'],
        ['Donnees', 'Base de Donnees', 'PostgreSQL (sur site, centre de donnees CAMPOST)'],
        ['Donnees', 'Stockage des Enregistrements', 'NAS/SAN local avec chiffrement'],
        ['Donnees', 'Moteur d\'Analyse', 'Intelligence d\'affaires et rapports'],
    ],
    col_widths=[3.5, 5, 7.5]
)

doc.add_heading('5.2 Principes de Conception', level=2)
add_bullet(doc, 'Sur site d\'abord : Tous les composants principaux heberges sur le centre de donnees CAMPOST pour la souverainete des donnees')
add_bullet(doc, 'Resilient a la faible bande passante : Optimise pour les connexions de 3-5 Mbps dans les bureaux ruraux, mode hors ligne disponible')
add_bullet(doc, 'Agents mobiles : Application web responsive permettant aux agents d\'utiliser des tablettes au lieu d\'ordinateurs complets')
add_bullet(doc, 'Bilingue par defaut : Toutes les interfaces, menus vocaux et rapports disponibles en francais et en anglais')
add_bullet(doc, 'Integration modulaire : Chaque systeme CAMPOST connecte via des connecteurs API standardises')

doc.add_page_break()

# ============================================================
# 6. COMPARAISON DES SOLUTIONS
# ============================================================
doc.add_heading('6. Comparaison des Solutions', level=1)
doc.add_paragraph(
    'Nous avons evalue 21 solutions a travers plusieurs categories. Voici une comparaison '
    'des options les plus pertinentes pour le contexte de la CAMPOST.'
)

doc.add_heading('6.1 Plateformes Cloud Entreprise', level=2)
add_table(doc,
    ['Solution', 'Points Forts', 'Faiblesses pour la CAMPOST', 'Cout Annuel Estime'],
    [
        ['Genesys Cloud CX', 'Leader du marche, multicanal complet, routage IA', 'Pas de centre de donnees africain, cout eleve, depend d\'internet', '90-180M FCFA+'],
        ['Five9', 'IA puissante, optimisation des effectifs', 'Centre sur les USA, pas de presence africaine, besoin de haut debit', '72-150M FCFA+'],
        ['NICE CXone', 'Centre de donnees africain (Afrique du Sud), niveau entreprise', 'Cout eleve, deploiement complexe', '78-170M FCFA+'],
        ['Talkdesk', 'Alimente par l\'IA, deploiements gouvernementaux', 'Pas de presence africaine, 100% cloud', '60-120M FCFA+'],
    ],
    col_widths=[3, 4.5, 5, 3]
)
doc.add_paragraph('Evaluation : Les plateformes cloud entreprise offrent des fonctionnalites puissantes mais sont couteuses, '
                   'dependantes d\'internet et sans presence locale au Cameroun. Non recommandees comme solution principale.')

doc.add_heading('6.2 Solutions Africaines', level=2)
add_table(doc,
    ['Solution', 'Points Forts', 'Faiblesses pour la CAMPOST', 'Cout Annuel Estime'],
    [
        ['PressOne Africa', 'Fonctionne a 3 Mbps, natif africain, echelle prouvee', 'Centre sur le Nigeria, presence francophone limitee', '12-36M FCFA'],
        ['HelloDuty', 'Construit en Afrique, SVI sans code, riche en API', 'Centre sur l\'Afrique de l\'Est, pas de marches francophones', '9-30M FCFA'],
        ['Ameyo (Exotel)', 'Cloud + sur site, support francais, presence africaine', 'Support specifique au Cameroun limite', '18-48M FCFA'],
        ['HoduCC', 'Abordable, deploiements au Nigeria, multicanal', 'Support base en Inde, decalage horaire', '9-27M FCFA'],
    ],
    col_widths=[3, 4.5, 5, 3]
)
doc.add_paragraph('Evaluation : Les solutions africaines offrent de meilleurs tarifs et une optimisation de la bande passante, '
                   'mais aucune n\'a de presence etablie au Cameroun/Afrique centrale.')

doc.add_heading('6.3 Open Source / Sur Site', level=2)
add_table(doc,
    ['Solution', 'Points Forts', 'Faiblesses pour la CAMPOST', 'Cout Estime'],
    [
        ['Asterisk', 'Gratuit, tres personnalisable, grande communaute', 'Necessite une expertise approfondie, pas d\'interface par defaut', '18-36M FCFA (services)'],
        ['Issabel', 'Gratuit, interface moderne, communaute francophone, PBX+CC complet', 'Communaute plus petite qu\'Asterisk', '15-30M FCFA (services)'],
        ['Vicidial', 'Gratuit, fonctionnalites de centre d\'appels eprouvees', 'Interface datee, oriente appels sortants', '12-27M FCFA (services)'],
        ['3CX', 'Option sur site, moderne, facile a utiliser', 'Fonctionnalites de centre d\'appels limitees, cout de licence', '6-18M FCFA (licence+config)'],
    ],
    col_widths=[3, 4.5, 5, 3.5]
)
doc.add_paragraph('Evaluation : Les solutions open source offrent le meilleur rapport cout et peuvent etre deployees sur le '
                   'centre de donnees de la CAMPOST. Necessitent un developpement personnalise pour les integrations CAMPOST.')

doc.add_heading('6.4 Specialisees et Postales', level=2)
add_table(doc,
    ['Solution', 'Points Forts', 'Faiblesses pour la CAMPOST', 'Cout Estime'],
    [
        ['UPU PTC Solutions', 'Integration postale native, compatible IPS/IFS', 'Fonctionnalites de centre d\'appels limitees', 'Frais d\'adhesion UPU'],
        ['Odoo VoIP+Helpdesk', 'ERP complet, open source, support francais', 'Module VoIP basique, necessite personnalisation', '12-30M FCFA'],
        ['Contaque', 'Appels bases sur GSM, fonctionnalites IA', 'Petit fournisseur, references africaines limitees', '15-36M FCFA'],
        ['Solution Sur Mesure', 'Ajustement parfait, controle total, souverainete des donnees', 'Cout initial plus eleve, delai plus long', '48-90M FCFA (dev)'],
    ],
    col_widths=[3, 4.5, 5, 3]
)

doc.add_page_break()

# ============================================================
# 7. APPROCHE RECOMMANDEE
# ============================================================
doc.add_heading('7. Approche Recommandee', level=1)
doc.add_paragraph(
    'Sur la base de notre evaluation approfondie, nous recommandons les approches suivantes, '
    'classees par ordre de pertinence pour le contexte specifique de la CAMPOST :'
)

doc.add_heading('7.1 Option A : Solution Hybride Sur Mesure (Recommandee)', level=2)
p = doc.add_paragraph()
run = p.add_run('Meilleur equilibre entre cout, controle, conformite et integration CAMPOST.')
run.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x6B, 0x1B)

doc.add_paragraph('Composants principaux :')
add_bullet(doc, 'Issabel/Asterisk : Central telephonique et moteur de centre d\'appels open source, deploye sur le centre de donnees CAMPOST')
add_bullet(doc, 'Poste de Travail Agent Personnalise : Application web avec vues integrees CampostMoney, UPU IPS et NPSI')
add_bullet(doc, 'API WhatsApp Business : Via CM.com ou integration directe Meta pour le multicanal')
add_bullet(doc, 'Passerelles GSM : Connexion directe aux reseaux MTN, Orange et CAMTEL')
add_bullet(doc, 'PostgreSQL + Stockage Local : Toutes les donnees restent dans le centre de donnees CAMPOST')

doc.add_paragraph()
add_table(doc,
    ['Aspect', 'Details'],
    [
        ['Budget Estime', 'FCFA 50-90 millions pour le developpement + Annee 1'],
        ['Cout Annuel de Fonctionnement', 'FCFA 12-25 millions pour la maintenance, telephonie, hebergement'],
        ['Delai de Mise en Oeuvre', '4-6 mois (deploiement progressif)'],
        ['Souverainete des Donnees', '100% sur site au centre de donnees CAMPOST'],
        ['Evolutivite', 'Nombre d\'agents illimite, evolution selon les besoins'],
        ['Conformite', 'Controle total sur la conformite a la Loi N. 2024/017'],
    ],
    col_widths=[5, 11]
)

doc.add_heading('7.2 Option B : Cloud Africain + Integration Personnalisee', level=2)
doc.add_paragraph('Utiliser PressOne Africa ou Ameyo comme moteur de centre d\'appels, avec des '
                   'connecteurs API personnalises pour les systemes CAMPOST.')
add_table(doc,
    ['Aspect', 'Details'],
    [
        ['Budget Estime', 'FCFA 30-55 millions pour la mise en place + Annee 1'],
        ['Cout Annuel de Fonctionnement', 'FCFA 18-45 millions abonnement + telephonie'],
        ['Delai de Mise en Oeuvre', '2-4 mois'],
        ['Souverainete des Donnees', 'Partielle - voix/agent sur le cloud, integrations sur site'],
        ['Risque', 'Dependance au fournisseur, couts d\'abonnement recurrents, questions de residence des donnees'],
    ],
    col_widths=[5, 11]
)

doc.add_heading('7.3 Option C : Plateforme Cloud Entreprise', level=2)
doc.add_paragraph('Deployer Genesys Cloud CX ou NICE CXone pour des fonctionnalites maximales.')
add_table(doc,
    ['Aspect', 'Details'],
    [
        ['Budget Estime', 'FCFA 90-180 millions pour l\'Annee 1'],
        ['Cout Annuel de Fonctionnement', 'FCFA 75-170 millions d\'abonnement'],
        ['Delai de Mise en Oeuvre', '3-6 mois'],
        ['Souverainete des Donnees', 'Limitee - donnees aux USA/Europe/Afrique du Sud'],
        ['Risque', 'Cout tres eleve, dependance totale a internet, support local limite'],
    ],
    col_widths=[5, 11]
)

doc.add_page_break()

# ============================================================
# 8. PLAN DE MISE EN OEUVRE
# ============================================================
doc.add_heading('8. Plan de Mise en Oeuvre (Option A)', level=1)

add_table(doc,
    ['Phase', 'Delai', 'Activites', 'Livrables'],
    [
        ['Phase 1 : Fondation', 'Semaines 1-4', 'Validation des exigences, mise en place de l\'infrastructure sur le centre de donnees CAMPOST, installation Issabel/Asterisk, configuration des passerelles GSM', 'Central telephonique fonctionnel, SVI basique (FR/EN), appels de test internes'],
        ['Phase 2 : Fonctionnalites', 'Semaines 5-10', 'Developpement du poste de travail agent, regles de routage, enregistrement des appels, systeme de tickets, tableau de bord superviseur', 'Centre d\'appels fonctionnel avec 10-20 agents, rapports basiques'],
        ['Phase 3 : Integrations', 'Semaines 11-16', 'Integration API CampostMoney, integration suivi UPU IPS, pont USSD NPSI #237#, canal WhatsApp Business', 'Centre multicanal avec integrations systemes CAMPOST'],
        ['Phase 4 : Pilote', 'Semaines 17-20', 'Pilote au siege de Yaounde + 5 bureaux de poste majeurs, formation des agents (FR/EN), optimisation des performances, tests d\'acceptation', 'Systeme valide, agents formes, indicateurs de performance'],
        ['Phase 5 : Deploiement', 'Semaines 21-24', 'Deploiement progressif vers les autres bureaux, mise en place des agents distants, alimentation de la base de connaissances, support de mise en production', 'Deploiement complet en production, documentation, transfert de competences'],
    ],
    col_widths=[3, 2.5, 5.5, 5]
)

doc.add_heading('8.1 Apres le Deploiement', level=2)
add_bullet(doc, 'Mois 7-12 : Integration de la plateforme e-commerce Bolamba lors de son lancement')
add_bullet(doc, 'En continu : Amelioration du SVI par intelligence artificielle, analyse des sentiments, routage predictif')
add_bullet(doc, 'Trimestriel : Revues de performance, recyclage des agents, mises a jour du systeme')

doc.add_page_break()

# ============================================================
# 9. CONFORMITE ET SECURITE
# ============================================================
doc.add_heading('9. Conformite et Securite', level=1)

doc.add_heading('9.1 Loi Camerounaise sur la Protection des Donnees (N. 2024/017)', level=2)
doc.add_paragraph('En vigueur le 23 juin 2026, cette loi exige :')
add_bullet(doc, 'Residence des donnees : Donnees clients stockees au Cameroun (centre de donnees CAMPOST)')
add_bullet(doc, 'Gestion du consentement : Consentement pour l\'enregistrement des appels')
add_bullet(doc, 'Notification de violation : Procedures de reponse aux incidents')
add_bullet(doc, 'Minimisation des donnees : Collecte uniquement des informations clients necessaires')
add_bullet(doc, 'Droit d\'acces/suppression : Portail de gestion des donnees clients')
doc.add_paragraph('Notre approche recommandee sur site (Option A) assure une conformite totale par defaut.')

doc.add_heading('9.2 Reglementations Financieres CEMAC', level=2)
doc.add_paragraph('Le Reglement COBAC N. 01/20 exige :')
add_bullet(doc, 'Enregistrement des appels pour les interactions liees aux services financiers (CampostMoney)')
add_bullet(doc, 'Piste d\'audit pour toutes les resolutions de litiges clients')
add_bullet(doc, 'Authentification securisee avant la divulgation d\'informations de compte')
add_bullet(doc, 'Rapports de conformite reguliers')

doc.add_heading('9.3 Mesures de Securite', level=2)
add_bullet(doc, 'Chiffrement de bout en bout pour les canaux vocaux et de donnees')
add_bullet(doc, 'Controle d\'acces base sur les roles (RBAC) pour les roles agent/superviseur/administrateur')
add_bullet(doc, 'Stockage chiffre des enregistrements d\'appels avec politiques de retention')
add_bullet(doc, 'Audits de securite reguliers et tests de penetration')
add_bullet(doc, 'Sauvegarde et reprise apres sinistre sur l\'infrastructure CAMPOST')

doc.add_page_break()

# ============================================================
# 10. ESTIMATION DES COUTS
# ============================================================
doc.add_heading('10. Estimation des Couts (Option A - Recommandee)', level=1)

doc.add_heading('10.1 Investissement Initial (Annee 1)', level=2)
add_table(doc,
    ['Poste', 'Description', 'Cout Estime (FCFA)'],
    [
        ['Mise en Place Infrastructure', 'Configuration serveur sur le DC CAMPOST, reseau, passerelles GSM', '8 000 000'],
        ['Deploiement Issabel/Asterisk', 'Installation du central, configuration SVI, lignes SIP', '5 000 000'],
        ['Developpement Poste Agent', 'Application web personnalisee avec integrations CAMPOST', '20 000 000'],
        ['Integration CampostMoney', 'Connecteur API pour les requetes de services financiers', '8 000 000'],
        ['Integration UPU IPS', 'Connecteur du systeme de suivi', '5 000 000'],
        ['Pont NPSI #237#', 'Integration du libre-service USSD', '6 000 000'],
        ['Canal WhatsApp', 'Configuration et integration de l\'API WhatsApp Business', '4 000 000'],
        ['Formation et Documentation', 'Formation des agents (FR/EN), guides administrateurs, videos', '4 000 000'],
        ['Materiel Passerelles GSM', 'Passerelles 4-8 ports pour MTN/Orange/CAMTEL', '3 000 000'],
        ['Tests et Assurance Qualite', 'Tests d\'acceptation, tests de charge, audit de securite', '3 000 000'],
        ['Gestion de Projet', 'Chef de projet, coordination, rapports (6 mois)', '6 000 000'],
        ['', 'TOTAL INVESTISSEMENT INITIAL', '72 000 000'],
    ],
    col_widths=[4.5, 7, 4]
)

doc.add_heading('10.2 Couts Recurrents Annuels', level=2)
add_table(doc,
    ['Poste', 'Description', 'Cout Annuel (FCFA)'],
    [
        ['Telephonie (SIP/GSM)', 'Minutes de voix, SMS, frais de lignes', '6 000 000'],
        ['API WhatsApp Business', 'Frais de messages et couts API', '2 000 000'],
        ['Maintenance et Support', 'Corrections, mises a jour, support 8x5', '8 000 000'],
        ['Hebergement (DC CAMPOST)', 'Electricite, espace rack, bande passante', '2 000 000'],
        ['Licences (le cas echeant)', '3CX ou modules supplementaires', '1 500 000'],
        ['Formation (continue)', 'Integration des nouveaux agents, recyclage', '2 000 000'],
        ['', 'TOTAL RECURRENT ANNUEL', '21 500 000'],
    ],
    col_widths=[4.5, 7, 4]
)

doc.add_heading('10.3 Cout Total de Possession sur 3 Ans', level=2)
add_table(doc,
    ['Periode', 'Cout (FCFA)'],
    [
        ['Annee 1 (Construction + Fonctionnement)', '93 500 000'],
        ['Annee 2 (Fonctionnement + Bolamba)', '28 000 000'],
        ['Annee 3 (Fonctionnement + Optimisation)', '21 500 000'],
        ['Cout Total sur 3 Ans', '143 000 000'],
    ],
    col_widths=[8, 5]
)

doc.add_page_break()

# ============================================================
# 11. POURQUOI NOUS CHOISIR
# ============================================================
doc.add_heading('11. Pourquoi Nous Choisir', level=1)
add_bullet(doc, 'Expertise Locale : Connaissance approfondie du paysage des telecommunications camerounaises (MTN, Orange, CAMTEL), du marche bilingue FR/EN et de l\'environnement commercial en FCFA.')
add_bullet(doc, 'Connaissance de l\'Ecosysteme CAMPOST : Recherche approfondie de l\'ecosysteme CAMPOST incluant CampostMoney, NPSI #237#, UPU IPS et la future plateforme Bolamba.')
add_bullet(doc, 'Approche Axee sur la Conformite : Solution concue des le depart pour la conformite a la Loi N. 2024/017 et aux reglements COBAC.')
add_bullet(doc, 'Technologie Rentable : Base open source eliminant les licences couteuses tout en offrant des fonctionnalites de niveau entreprise.')
add_bullet(doc, 'Souverainete des Donnees : Deploiement 100% sur site sur le propre centre de donnees de la CAMPOST - aucune donnee ne quitte le Cameroun.')
add_bullet(doc, 'Evolutif et Pret pour l\'Avenir : Architecture modulaire prete pour l\'integration de Bolamba et les ameliorations par intelligence artificielle.')

doc.add_page_break()

# ============================================================
# 12. CONCLUSION
# ============================================================
doc.add_heading('12. Conclusion', level=1)
doc.add_paragraph(
    'La CAMPOST se trouve a un moment charniere de sa transformation numerique. Avec '
    'CampostMoney desservant une base de clients croissante, la plateforme NPSI #237# '
    'connectant les banques a l\'echelle nationale et la plateforme e-commerce Bolamba '
    'qui sera lancee en 2026, un centre d\'appels moderne n\'est plus une option - c\'est '
    'une infrastructure essentielle pour la satisfaction client et la qualite de service.'
)
doc.add_paragraph(
    'Notre approche hybride recommandee (Option A) fournit une solution de centre d\'appels '
    'complete et integree a la CAMPOST a une fraction du cout des plateformes cloud '
    'entreprise, tout en garantissant la souverainete totale des donnees et la conformite '
    'reglementaire. En tirant parti du centre de donnees existant de 52 millions de dollars '
    'de la CAMPOST et de technologies open source eprouvees, nous minimisons les couts '
    'recurrents et maximisons la valeur a long terme.'
)
doc.add_paragraph(
    'Le cout total de possession sur 3 ans d\'environ 143 millions de FCFA se compare '
    'favorablement aux alternatives entreprise qui couteraient entre 225 et 510 millions de '
    'FCFA sur la meme periode, tout en offrant une integration superieure avec l\'ecosysteme '
    'unique de la CAMPOST.'
)

p = doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Nous nous rejouissons a l\'idee de nous associer a la CAMPOST pour ce projet transformateur.')
run.bold = True
run.italic = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

doc.add_page_break()

# ============================================================
# ANNEXE : SOURCES
# ============================================================
doc.add_heading('Annexe : Sources de Recherche', level=1)
sources = [
    'Campost mise 3 milliards FCFA sur Bolamba pour le e-commerce - Business in Cameroon',
    'Partenariat Campost + FindMe pour l\'adressage numerique - We Are Tech Africa',
    'Plateforme bancaire numerique CampostMoney - campostmoney.com',
    'Implementation CampostMoney par Sixgoldtech - sixgoldtech.com',
    'Lancement de la plateforme de paiement USSD #237# - Digital Business Africa',
    'Le Cameroun premier pays africain avec un code USSD national - SUPTIC',
    'Partenariat Camwater-Campost pour le paiement de factures - Investir au Cameroun',
    'Centre de donnees Campost (AS328123) - PeeringDB',
    'Le secteur telecom camerounais croit de 18% en 2024 - TechAfrica News',
    'Rapport sur les operateurs telecom du Cameroun 2025 - GlobeNewsWire',
    'Loi camerounaise sur la protection des donnees N. 2024/017 - LEX Africa',
    'Cadre de protection des donnees CEMAC - DataGuidance',
    'Instance NICE CXone en Afrique du Sud (Dec 2025) - NICE',
    'Meilleurs fournisseurs de solutions de centres d\'appels en Afrique (2024) - Telebu',
    'Centre d\'appels PressOne Africa - pressone.africa',
    'Centre d\'appels HelloDuty avec IA - helloduty.com',
    'Centre de contact cloud Ameyo en Afrique - ameyo.com',
    'Logiciel de centre de contact HoduCC - hodusoft.com',
    'Centre de Technologie Postale de l\'UPU - upu.int',
    'Modernisation gouvernementale MINPOSTEL - minpostel.gov.cm',
]
for i, src in enumerate(sources, 1):
    p = doc.add_paragraph(f'{i}. {src}')
    p.runs[0].font.size = Pt(10)

output_path = '/Users/user/Documents/hairstyle/hairstyle/campost-call-center-software/Proposition_Centre_Appels_CAMPOST.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
